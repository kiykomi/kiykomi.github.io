#!/usr/bin/env python3
"""Translate Kyiko.M_CV.docx to EN/RU/UK while preserving run formatting."""

from __future__ import annotations

import json
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Kyiko.M_CV.docx"
TRANSLATIONS_FILE = ROOT / "scripts" / "cv_translations.json"

OUTPUTS = {
    "en": ROOT / "Kyiko.M_CV_EN.docx",
    "ru": ROOT / "Kyiko.M_CV_RU.docx",
    "uk": ROOT / "Kyiko.M_CV_UK.docx",
}


def run_group_key(run) -> tuple:
    return (run.bold, run.italic, run.font.name, str(run.font.size))


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def grouped_runs(paragraph):
    groups: list[dict] = []
    for run in paragraph.runs:
        text = run.text
        if not text and not groups:
            continue
        key = run_group_key(run)
        if not groups or groups[-1]["key"] != key:
            groups.append({"key": key, "runs": [run], "text": text})
        else:
            groups[-1]["runs"].append(run)
            groups[-1]["text"] += text
    return groups


def apply_translation(paragraph, translations: dict, lang: str) -> list[str]:
    missing: list[str] = []
    for group in grouped_runs(paragraph):
        source = group["text"]
        if not source.strip():
            continue
        entry = translations.get(source)
        if entry is None:
            missing.append(source)
            continue
        target = entry.get(lang)
        if target is None:
            missing.append(source)
            continue
        group["runs"][0].text = target
        for run in group["runs"][1:]:
            run.text = ""
    return missing


def translate_doc(source: Path, target: Path, translations: dict, lang: str) -> list[str]:
    shutil.copy2(source, target)
    doc = Document(target)
    missing: list[str] = []
    for paragraph in iter_all_paragraphs(doc):
        missing.extend(apply_translation(paragraph, translations, lang))
    doc.save(target)
    return sorted(set(missing))


def main() -> None:
    translations = json.loads(TRANSLATIONS_FILE.read_text(encoding="utf-8"))
    for lang, output in OUTPUTS.items():
        missing = translate_doc(SOURCE, output, translations, lang)
        print(f"{output.name}: saved")
        if missing:
            print(f"  missing ({len(missing)}):")
            for item in missing:
                print(f"    - {item[:120]!r}")


if __name__ == "__main__":
    main()
